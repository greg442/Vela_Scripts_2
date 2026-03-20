#!/usr/bin/env python3
"""
Agent Delivery Script — Hannah CoS System
Usage:
  python3 deliver_report.py \
    --title "COLEL Marketing Strategy: 65+ HNWI Segment" \
    --agent marketing \
    --filename "2026-03-12-colel-65plus-strategy" \
    --content-file /tmp/report_content.json

content-file is a JSON file with this structure:
{
  "sections": [
    {"heading": "1. Positioning", "body": "..."},
    {"heading": "2. Key Messaging Themes", "body": "..."},
    ...
  ],
  "summary_bullets": ["bullet 1", "bullet 2", "bullet 3"]
}
"""

import argparse
import json
import os
import sys
import requests

AGENT_TOPICS = {
    "marketing": "76",
    "analyst": "77",
    "researcher": "78",
    "pm": "79",
    "legal": "80",
}

AGENT_DRIVE_FOLDERS = {
    "marketing": "1S0WkBC4PXTW6O3AEIzlCX7Y66Yw7zNqW",
    "analyst":   "1t57G63mJmVVx3te9Bwdea2ABfx_u_rqy",
    "researcher":"1Ylqig0q3Y_eAM8jsFwOfpGwLrYvmAv0T",
    "pm":        "1ZFISkXUbyGjTEFrwZzBkTIHuRfYobMp-",
    "legal":     "1Q0LDvGGKd5nv-s1gQnrsM_HPPgjMAxW1",
}

GOG_ACCOUNT = "greg@gregshindler.com"
DRIVE_AUTH_PATH = "/Users/gregshindler/.openclaw/workspace-cos/scripts/.drive_auth.json"

CHAT_ID = "-1003750313044"
BOT_TOKEN = "8459904439:AAE4r3u7lyyQ6E_DiataGN40Nzv_6svH0ug"
OUTPUT_BASE = "/Users/gregshindler/.openclaw/workspace-cos/reference"


def get_drive_access_token():
    """Get a Google Drive access token from local auth config (no keychain, no prompts)."""
    import json as _json

    try:
        with open(DRIVE_AUTH_PATH) as f:
            auth = _json.load(f)

        resp = requests.post("https://oauth2.googleapis.com/token", data={
            "client_id": auth["client_id"],
            "client_secret": auth["client_secret"],
            "refresh_token": auth["refresh_token"],
            "grant_type": "refresh_token"
        })
        return resp.json().get("access_token")
    except Exception as e:
        print(f"⚠️  Token error: {e}")
        return None


def create_docx(title, filename, agent, sections):
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        os.system("pip3 install python-docx -q --break-system-packages")
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    t = doc.add_heading(title, 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    from datetime import date
    s = doc.add_paragraph(f"Prepared by: {agent.title()} Agent  |  {date.today().strftime('%B %d, %Y')}")
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    for section in sections:
        doc.add_heading(section["heading"], level=1)
        doc.add_paragraph(section["body"])

    out_dir = os.path.join(OUTPUT_BASE, agent)
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, f"{filename}.docx")
    doc.save(out_path)
    return out_path


def post_to_telegram(agent, title, filepath, bullets):
    thread_id = AGENT_TOPICS.get(agent.lower())
    if not thread_id:
        print(f"ERROR: Unknown agent '{agent}'. Valid: {list(AGENT_TOPICS.keys())}")
        sys.exit(1)

    rel_path = filepath.replace("/Users/gregshindler/.openclaw/workspace-cos/", "")
    bullet_text = "\n".join(f"• {b}" for b in bullets)
    message = f"📄 *{title}*\n\nFile: `{rel_path}`\n\n{bullet_text}"

    url = f"https://api.telegram.org/bot{BOT_TOKEN}/sendMessage"
    resp = requests.post(url, json={
        "chat_id": CHAT_ID,
        "message_thread_id": int(thread_id),
        "text": message,
        "parse_mode": "Markdown"
    })
    data = resp.json()
    if data.get("ok"):
        print(f"✅ Posted to {agent.title()} topic (thread {thread_id})")
    else:
        print(f"❌ Telegram error: {data}")
        sys.exit(1)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--title", required=True)
    parser.add_argument("--agent", required=True, choices=list(AGENT_TOPICS.keys()))
    parser.add_argument("--filename", required=True)
    parser.add_argument("--content-file", required=True)
    args = parser.parse_args()

    with open(args.content_file) as f:
        content = json.load(f)

    sections = content.get("sections", [])
    bullets = content.get("summary_bullets", ["See full document for details."])

    if not sections:
        print("ERROR: No sections found in content file.")
        sys.exit(1)

    print(f"Creating Word doc: {args.filename}.docx ...")
    docx_path = create_docx(args.title, args.filename, args.agent, sections)
    print(f"Saved: {docx_path}")

    print(f"Uploading to Google Drive ...")
    folder_id = AGENT_DRIVE_FOLDERS.get(args.agent)
    if folder_id:
        access_token = get_drive_access_token()
        if access_token:
            result = os.popen(
                f'gog drive upload "{docx_path}" --access-token "{access_token}" --parent {folder_id} --json 2>&1'
            ).read()
            if '"id"' in result:
                print(f"✅ Uploaded to Drive ({args.agent.title()} folder)")
            else:
                print(f"⚠️  Drive upload issue: {result.strip()}")
        else:
            print("⚠️  Could not obtain Drive access token")
    else:
        print("⚠️  No Drive folder configured for this agent")

    print(f"Posting to Telegram topic: {args.agent} ...")
    post_to_telegram(args.agent, args.title, docx_path, bullets)


if __name__ == "__main__":
    main()
